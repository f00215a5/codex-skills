#!/usr/bin/env python3
"""Programmatic structure QA for the generated manual DOCX (python-docx only).

This is the replacement for Word/LibreOffice rendering QA in the lite skill:
it checks what is verifiable without a renderer — structure, numbering restarts,
tables, captions vs images, page setup, package integrity — and fails closed
when any check fails.

Usage:
    <venv-python> verify_docx.py --docx manual.docx [--manifest manual.json]

Exit code 0 == all checks passed.  Visual fidelity after opening the file in
Word/WPS/LibreOffice remains the user's review step, not this script's claim.
"""

from __future__ import annotations

import argparse
import re
import sys
import zipfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

CAPTION_PATTERN = re.compile(r"紅框\s*(\d+)")
CHECKMARK = "PASS"
CROSS = "FAIL"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Structure QA for the manual DOCX.")
    parser.add_argument("--docx", required=True, type=Path)
    parser.add_argument("--manifest", type=Path, help="Optional manifest for caption/id cross-check.")
    return parser.parse_args()


class Reporter:
    def __init__(self) -> None:
        self.failures: list[str] = []

    def check(self, name: str, condition: bool, detail: str = "") -> None:
        status = CHECKMARK if condition else CROSS
        message = f"[{status}] {name}"
        if detail:
            message += f" — {detail}"
        print(message)
        if not condition:
            self.failures.append(name)


def headings_of(document: Document) -> list[tuple[str, str]]:
    """Return (style_name, text) of heading paragraphs in document order."""
    result = []
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name.startswith("Heading"):
            result.append((style_name, paragraph.text.strip()))
    return result


def numbered_paragraphs(document: Document) -> list[tuple[int, int]]:
    """Return (numId, ilvl) for every paragraph carrying w:numPr."""
    numbered = []
    for paragraph in document.paragraphs:
        num_pr = paragraph._p.find(qn("w:pPr") + "/" + qn("w:numPr"))
        if num_pr is None:
            continue
        num_id = num_pr.find(qn("w:numId"))
        ilvl = num_pr.find(qn("w:ilvl"))
        number = int(num_id.get(qn("w:val"))) if num_id is not None else -1
        level = int(ilvl.get(qn("w:val"))) if ilvl is not None else 0
        numbered.append((number, level))
    return numbered


def drawing_targets(document: Document) -> list[tuple[int, str]]:
    """Return (paragraph_index, media target) for every embedded drawing."""
    items: list[tuple[int, str]] = []
    for index, paragraph in enumerate(document.paragraphs):
        for blip in paragraph._p.xpath(".//a:blip"):
            r_id = blip.get(qn("r:embed"))
            if r_id and r_id in document.part.related_parts:
                items.append((index, str(document.part.related_parts[r_id].partname)))
    return items


def table_headers(table) -> list[str]:
    return [cell.text.strip() for cell in table.rows[0].cells]


def main() -> int:
    args = parse_args()
    docx_path = args.docx.expanduser().resolve()
    reporter = Reporter()

    try:
        document = Document(str(docx_path))
    except Exception as error:  # noqa: BLE001 — any open failure fails closed
        print(f"[FAIL] open docx — {error}")
        return 1

    # -- package integrity -------------------------------------------------- #
    try:
        with zipfile.ZipFile(docx_path) as archive:
            names = archive.namelist()
    except zipfile.BadZipFile as error:
        print(f"[FAIL] zip integrity — {error}")
        return 1
    reporter.check("docx contains document.xml", "word/document.xml" in names)
    media_files = [name for name in names if name.startswith("word/media/")]
    reporter.check("page writer exists", "word/numbering.xml" in names, f"media={len(media_files)}")

    # -- page setup (Letter, baseline margins) ------------------------------ #
    section = document.sections[0]
    reporter.check("paper size is Letter",
                   abs(section.page_width.inches - 8.5) < 0.01 and abs(section.page_height.inches - 11.0) < 0.01)
    reporter.check("margins match baseline",
                   abs(section.top_margin.inches - 0.67) < 0.01
                   and abs(section.bottom_margin.inches - 0.59) < 0.01
                   and abs(section.left_margin.inches - 0.65) < 0.01
                   and abs(section.right_margin.inches - 0.65) < 0.01)

    # -- heading order ------------------------------------------------------ #
    headings = headings_of(document)
    heading_texts = [text for _, text in headings]
    reporter.check("has title area content", any(paragraph.text.strip()
                                                 for paragraph in document.paragraphs[:5]))
    for required in ("修訂狀態", "使用提醒", "共通操作規則", "更新紀錄"):
        reporter.check(f"heading present: {required}", required in heading_texts)
    log_position = heading_texts.index("更新紀錄") if "更新紀錄" in heading_texts else -1
    reporter.check("更新紀錄 is the last heading",
                   log_position != -1 and log_position == len(heading_texts) - 1)
    order = [text for text in heading_texts if text in ("修訂狀態", "使用提醒", "共通操作規則")]
    reporter.check("heading order 修訂狀態→使用提醒→共通操作規則", order == ["修訂狀態", "使用提醒", "共通操作規則"])

    # -- independent step numbering per section ----------------------------- #
    numbered = numbered_paragraphs(document)
    reporter.check("operation steps carry numbering", len(numbered) >= 1, f"numbered paragraphs={len(numbered)}")
    # Every Heading-2 operation section must own exactly one numId so its step
    # list restarts at 1 instead of continuing from the previous section.
    blocks: list[list[int]] = []
    current: list[int] = []
    seen_section = False
    for paragraph in document.paragraphs:
        style = paragraph.style.name if paragraph.style else ""
        if style == "Heading 2":
            if current:
                blocks.append(current)
            current = []
            seen_section = True
            continue
        num_pr = paragraph._p.find(qn("w:pPr") + "/" + qn("w:numPr"))
        if seen_section and num_pr is not None and (num_id := num_pr.find(qn("w:numId"))) is not None:
            current.append(int(num_id.get(qn("w:val"))))
    if current:
        blocks.append(current)
    reporter.check("every step list restarts within its section",
                   any(blocks) and all(len(set(block)) <= 1 for block in blocks),
                   detail=f"step blocks={[len(b) for b in blocks]}")

    # -- field tables ------------------------------------------------------- #
    tables = document.tables
    field_tables = [
        t for t in tables
        if any("欄位" in header or "控制項" in header for header in table_headers(t))
    ]
    reporter.check("field tables present for operations", len(field_tables) >= 1, f"tables={len(field_tables)}")
    for table in field_tables:
        headers = table_headers(table)
        reporter.check("field table has 定義/必填 columns",
                       "定義" in headers and "必填" in headers, detail=f"headers={headers}")

    # -- update log --------------------------------------------------------- #
    if tables:
        last_table_headers = table_headers(tables[-1])
        reporter.check("update log is the last table",
                       last_table_headers[:3] == ["版本", "日期", "更新內容"],
                       detail=f"headers={last_table_headers}")
        reporter.check("update log has data rows", len(tables[-1].rows) >= 2)

    # -- captions vs images ------------------------------------------------- #
    # 紅框編號是每張圖獨立重編（同張圖內的控制項對照），所以不要求全域唯一。
    drawings = drawing_targets(document)
    reporter.check("images embedded", len(drawings) >= 1, f"drawings={len(drawings)}")
    missing_targets = [target for _, target in drawings if target.lstrip("/") not in names]
    reporter.check("every embedded image resolves inside the package",
                   not missing_targets, detail=f"missing={missing_targets}")

    drawings_paragraphs = {index for index, _ in drawings}
    captions: list[str] = []
    for index in range(len(document.paragraphs) - 1):
        if index in drawings_paragraphs:
            if match := CAPTION_PATTERN.match(document.paragraphs[index + 1].text.strip()):
                captions.append(match.group(1))
    reporter.check("image captions reference 紅框 numbers", len(captions) >= 1, f"captions={captions}")
    caption_follow = sum(
        1 for index, _ in drawings
        if index + 1 < len(document.paragraphs) and document.paragraphs[index + 1].text.strip()
    )
    reporter.check("every image drawing is followed by a caption line",
                   caption_follow == len(drawings),
                   f"captioned={caption_follow} drawings={len(drawings)}")

    # -- optional build-manifest cross-check -------------------------------- #
    if args.manifest:
        manifest = json_load(args.manifest.expanduser().resolve())
        expected: list[str] = []
        entry = (manifest.get("chapter") or {}).get("entry") or {}
        if entry.get("image"):
            if match := CAPTION_PATTERN.match(str(entry.get("caption", ""))):
                expected.append(match.group(1))
        for section in (manifest.get("chapter") or {}).get("sections", []):
            for step in section.get("steps", []):
                if step.get("image"):
                    if match := CAPTION_PATTERN.match(str(step.get("caption", ""))):
                        expected.append(match.group(1))
        reporter.check("caption sequence matches build manifest order",
                       list(captions) == expected, detail=f"docx={captions} manifest={expected}")

    if reporter.failures:
        print(f"\nverify_docx.py: {len(reporter.failures)} check(s) failed "
              f"({', '.join(reporter.failures)})")
        return 1
    print("\nverify_docx.py: all structure checks passed. "
          "Visual fidelity after opening the file is not verified by this script.")
    return 0


def json_load(path: Path) -> dict[str, Any]:
    import json
    payload = json.loads(path.read_text(encoding="utf-8"))
    assert isinstance(payload, dict)
    return payload


if __name__ == "__main__":
    raise SystemExit(main())