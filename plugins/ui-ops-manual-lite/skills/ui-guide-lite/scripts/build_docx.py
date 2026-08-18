#!/usr/bin/env python3
"""Build the UI manual DOCX from a structured manifest (Python-only).

The document is written deterministically and afterwards validated by
verify_docx.py. Runtime availability and visual-verification limitations are
reported in the conversation; this builder rejects those warnings as DOCX
content (see references/document-structure-qa.md).

Usage:
    <venv-python> build_docx.py --manifest manual.json --output manual.docx

Manifest schema (relative image paths resolve against the manifest's folder,
or against "baseDir" when present):

{
  "title": "訂單管理操作說明書",
  "subtitle": "系統後台操作指引",
  "applicableScreens": ["系統後台 > 訂單管理 > 訂單列表"],
  "revision": {"version": "1.0", "date": "2026-01-15", "summary": "初版"},
  "versionBasis": {"source": "...", "basis": "...", "date": "...",
                   "features": ["..."]},              # optional
  "usageReminders": ["...", "..."],
  "commonRules": ["...", "..."],
  "fontName": "PingFang TC",
  "chapter": {
    "title": "訂單管理功能",
    "entry": {"caption": "功能入口：...", "image": "annotated/entry.png",
              "detail": "..."},                       # detail optional
    "sections": [
      {
        "title": "查詢訂單",
        "preconditions": "...",                       # optional
        "steps": [
          {"caption": "紅框 1：點選「訂單查詢」。",     # caption is the 圖說
           "image": "annotated/step1.png",            # optional
           "detail": "..."}                           # optional
        ],
        "fields": [
          {"name": "訂單編號", "definition": "...", "required": "選填",
           "limits": "最大 30 字元", "display": "查詢區"}  # display optional
        ],
        "impact": "...",          # 修改成功後的影響
        "failure": "...",         # 失敗／取消行為 (optional)
        "verification": "..."     # 操作後檢核
      }
    ]
  },
  "updateLog": [{"version": "1.0", "date": "2026-01-15", "changes": "初版"}]
}

Rule: the 更新紀錄 table lists entries from oldest to newest, in manifest order.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

TEAL = RGBColor(0x0F, 0x5B, 0x5B)
BODY_FONT_SIZE = Pt(11)
DEFAULT_FONT = "Microsoft JhengHei"
PAGE_W, PAGE_H = 8.5, 11.0
MARGIN_LEFT = MARGIN_RIGHT = 0.65
MARGIN_TOP, MARGIN_BOTTOM = 0.67, 0.59
RUNTIME_TOOL_PATTERN = re.compile(
    r"(?<![a-z])(?:word-render|libreoffice|word)(?![a-z])", re.IGNORECASE
)
RUNTIME_AVAILABILITY_MARKERS = (
    "未安裝",
    "不可用",
    "無法使用",
    "無法執行",
    "not installed",
    "unavailable",
    "not available",
    "cannot use",
)

BULLET_ABSNUM = """<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="{abs_id}">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="bullet"/>
    <w:lvlText w:val="•"/><w:lvlJc w:val="left"/></w:lvl>
</w:abstractNum>"""

STEP_ABSNUM = """<w:abstractNum xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:abstractNumId="{abs_id}">
  <w:multiLevelType w:val="hybridMultilevel"/>
  <w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/><w:lvlText w:val="%1."/>
    <w:lvlJc w:val="left"/><w:suff w:val="space"/></w:lvl>
</w:abstractNum>"""

NUM_XML = """<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:numId="{num_id}">
  <w:abstractNumId w:val="{abs_id}"/>
</w:num>"""


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build the UI manual DOCX from a manifest.")
    parser.add_argument("--manifest", required=True, type=Path)
    parser.add_argument("--output", required=True, type=Path)
    return parser.parse_args()


def load_manifest(path: Path) -> dict[str, Any]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(payload, dict):
        raise ValueError(f"{path}: manifest must be a JSON object")
    for key in ("title", "revision", "chapter", "updateLog"):
        if key not in payload:
            raise ValueError(f"{path}: missing required key {key!r}")
    reject_runtime_availability_warnings(payload)
    return payload


def iter_manifest_strings(value: Any, path: str = "$") -> Iterable[tuple[str, str]]:
    if isinstance(value, str):
        yield path, value
    elif isinstance(value, dict):
        for key, nested in value.items():
            yield from iter_manifest_strings(nested, f"{path}.{key}")
    elif isinstance(value, list):
        for index, nested in enumerate(value):
            yield from iter_manifest_strings(nested, f"{path}[{index}]")


def reject_runtime_availability_warnings(manifest: dict[str, Any]) -> None:
    for path, text in iter_manifest_strings(manifest):
        normalized = re.sub(r"\s+", " ", text.casefold())
        mentions_tool = RUNTIME_TOOL_PATTERN.search(normalized) is not None
        mentions_availability = any(marker in normalized for marker in RUNTIME_AVAILABILITY_MARKERS)
        if mentions_tool and mentions_availability:
            raise ValueError(
                f"{path}: runtime availability warning must be reported in chat, not DOCX"
            )


def base_dir(manifest: dict[str, Any], manifest_path: Path) -> Path:
    return Path(manifest.get("baseDir", manifest_path.parent)).expanduser().resolve()


def resolve_image(manifest_ref: str, base: Path) -> Path:
    path = base / manifest_ref
    if not path.is_file():
        raise FileNotFoundError(f"image not found: {path}")
    return path


# --------------------------------------------------------------------------- #
# numbering & run helpers
# --------------------------------------------------------------------------- #

def _next_element_id(root, tag: str) -> int:
    ids = [int(node.get(qn("w:" + tag))) for node in root.findall(qn("w:" + tag))]
    return (max(ids) + 1) if ids else 0


def _numbering_root(document: DocumentType):
    numbering_part = document.part.numbering_part
    return numbering_part.element


def fresh_num_id(document: DocumentType, abstract_xml: str) -> int:
    """Register a brand-new numbering definition and return its numId.

    Each call owns its own abstractNum + num, so every operation section
    restarts its numbering at step 1 (documented requirement).
    """

    from docx.oxml import parse_xml
    root = _numbering_root(document)
    abs_id = _next_element_id(root, "abstractNumId")
    num_id = _next_element_id(root, "numId")
    root.append(parse_xml(abstract_xml.format(abs_id=abs_id)))
    root.append(parse_xml(NUM_XML.format(num_id=num_id, abs_id=abs_id)))
    return num_id


def apply_numbering(paragraph: Paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)


def set_run_fonts(run, font_name: str, size: Pt | None = None) -> None:
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "zh-TW")
    r_pr.append(lang)
    if size is not None:
        run.font.size = size


def shade_paragraph(paragraph: Paragraph, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    paragraph._p.get_or_add_pPr().append(shd)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


# --------------------------------------------------------------------------- #
# document construction
# --------------------------------------------------------------------------- #

def setup_page(document: DocumentType) -> None:
    section = document.sections[0]
    section.page_width = Inches(PAGE_W)
    section.page_height = Inches(PAGE_H)
    section.left_margin = Inches(MARGIN_LEFT)
    section.right_margin = Inches(MARGIN_RIGHT)
    section.top_margin = Inches(MARGIN_TOP)
    section.bottom_margin = Inches(MARGIN_BOTTOM)


def setup_styles(document: DocumentType, font_name: str, heading_color: RGBColor) -> None:
    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = BODY_FONT_SIZE
    r_pr = normal.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[name]
        style.font.color.rgb = heading_color
        style_r_pr = style.element.get_or_add_rPr()
        style_r_fonts = style_r_pr.get_or_add_rFonts()
        style_r_fonts.set(qn("w:ascii"), font_name)
        style_r_fonts.set(qn("w:hAnsi"), font_name)
        style_r_fonts.set(qn("w:eastAsia"), font_name)


def add_centered_line(document: DocumentType, text: str, size: float, color: RGBColor, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.color.rgb = color
    set_run_fonts(run, document.styles["Normal"].font.name, Pt(size))


def add_heading(document: DocumentType, text: str, level: int) -> Paragraph:
    return document.add_heading(text, level=level)


def add_body_paragraph(document: DocumentType, text: str, style: str = "Normal") -> Paragraph:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_fonts(run, document.styles["Normal"].font.name)
    return paragraph


def add_bullet(document: DocumentType, text: str, num_id: int) -> Paragraph:
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_fonts(run, document.styles["Normal"].font.name)
    apply_numbering(paragraph, num_id)
    return paragraph


def add_image_with_caption(document: DocumentType, image_path: Path, caption: str, *, width: float = 6.0) -> None:
    document.add_picture(str(image_path), width=Inches(width))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    set_run_fonts(run, document.styles["Normal"].font.name)


def heading_cell(cell, text: str, font_name: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = True
    set_run_fonts(run, font_name)
    shade_cell(cell, "CDEBE5")


def build_tables(document: DocumentType, rows: Iterable[list[str]], header: list[str],
                 font_name: str = DEFAULT_FONT) -> None:
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for index, title in enumerate(header):
        heading_cell(table.rows[0].cells[index], title, font_name)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            text = cells[index].paragraphs[0].add_run(str(value))
            set_run_fonts(text, font_name)


# --------------------------------------------------------------------------- #
# content blocks
# --------------------------------------------------------------------------- #

def add_title_block(document: DocumentType, manifest: dict[str, Any]) -> None:
    add_centered_line(document, manifest["title"], 26, TEAL, bold=True)
    add_centered_line(document, manifest.get("subtitle", ""), 14, TEAL)
    screens = "、".join(manifest.get("applicableScreens", []))
    if screens:
        add_centered_line(document, f"適用畫面：{screens}", 11, RGBColor(0x66, 0x66, 0x66))


def add_revision_status(document: DocumentType, revision: dict[str, Any]) -> None:
    add_heading(document, "修訂狀態", 1)
    build_tables(
        document,
        rows=[[revision.get("version", ""), revision.get("date", ""), revision.get("summary", "")]],
        header=["文件版本", "日期", "本次修訂摘要"],
        font_name=document.styles["Normal"].font.name,
    )


def add_version_basis(document: DocumentType, basis: dict[str, Any]) -> None:
    if not basis:
        return
    add_heading(document, "版本依據", 1)
    for label, key in (("來源", "source"), ("比較基準", "basis"), ("發佈時間", "date")):
        if basis.get(key):
            add_body_paragraph(document, f"{label}：{basis[key]}")
    features = basis.get("features") or []
    if features:
        features_paragraph = add_body_paragraph(document, "本次納入功能：" + "、".join(features))


def add_usage_reminders(document: DocumentType, reminders: list[str]) -> None:
    add_heading(document, "使用提醒", 1)
    paragraph = add_body_paragraph(document, reminders[0] if reminders else "")
    shade_paragraph(paragraph, "E8F4EF")
    for reminder in reminders[1:]:
        next_paragraph = document.add_paragraph()
        run = next_paragraph.add_run(reminder)
        set_run_fonts(run, document.styles["Normal"].font.name)
        shade_paragraph(next_paragraph, "E8F4EF")


def add_common_rules(document: DocumentType, rules: list[str]) -> None:
    add_heading(document, "共通操作規則", 1)
    bullet_num = fresh_num_id(document, BULLET_ABSNUM)
    for rule in rules:
        add_bullet(document, rule, bullet_num)


def add_chapter(document: DocumentType, manifest: dict[str, Any], base: Path, font_name: str) -> None:
    chapter = manifest["chapter"]
    add_heading(document, chapter["title"], 1)

    entry = chapter.get("entry") or {}
    image_ref = entry.get("image")
    if image_ref:
        add_image_with_caption(document, resolve_image(image_ref, base), entry.get("caption", ""))
    if entry.get("detail"):
        add_body_paragraph(document, entry["detail"])

    for section in chapter.get("sections", []):
        add_heading(document, section["title"], 2)
        if section.get("preconditions"):
            add_body_paragraph(document, "前置條件：" + section["preconditions"])

        step_num = fresh_num_id(document, STEP_ABSNUM)
        for step in section.get("steps", []):
            step_paragraph = document.add_paragraph()
            run = step_paragraph.add_run(step.get("caption", ""))
            run.bold = True
            set_run_fonts(run, font_name)
            apply_numbering(step_paragraph, step_num)
            if step.get("detail"):
                add_body_paragraph(document, step["detail"])
            image_ref = step.get("image")
            if image_ref:
                add_image_with_caption(
                    document, resolve_image(image_ref, base), step.get("caption", "")
                )

        fields = section.get("fields") or []
        if fields:
            header = ["欄位或控制項", "定義", "必填", "限制／選項"]
            has_display = any(field.get("display") for field in fields)
            if has_display:
                header.append("顯示條件／結果")
            rows = []
            for field in fields:
                row = [field.get("name", ""), field.get("definition", ""),
                       field.get("required", ""), field.get("limits", "")]
                if has_display:
                    row.append(field.get("display", ""))
                rows.append(row)
            build_tables(document, rows, header, font_name)

        for label, key in (("修改成功後的影響", "impact"),
                           ("失敗／取消行為", "failure"),
                           ("操作後檢核", "verification")):
            if section.get(key):
                add_body_paragraph(document, f"{label}：{section[key]}")


def add_update_log(document: DocumentType, update_log: list[dict[str, Any]]) -> None:
    document.add_page_break()
    add_heading(document, "更新紀錄", 1)
    rows = [[entry.get("version", ""), entry.get("date", ""), entry.get("changes", "")]
            for entry in update_log]
    build_tables(document, rows, header=["版本", "日期", "更新內容"],
                 font_name=document.styles["Normal"].font.name)


def build(document: DocumentType, manifest: dict[str, Any], base: Path) -> None:
    font_name = manifest.get("fontName", DEFAULT_FONT)
    setup_page(document)
    setup_styles(document, font_name, TEAL)
    add_title_block(document, manifest)
    add_revision_status(document, manifest["revision"])
    add_version_basis(document, manifest.get("versionBasis"))
    add_usage_reminders(document, manifest.get("usageReminders") or [])
    add_common_rules(document, manifest.get("commonRules") or [])
    add_chapter(document, manifest, base, font_name)
    add_update_log(document, manifest["updateLog"])


def main() -> int:
    args = parse_args()
    manifest_path = args.manifest.expanduser().resolve()
    output_path = args.output.expanduser().resolve()
    try:
        manifest = load_manifest(manifest_path)
        base = base_dir(manifest, manifest_path)
        document = Document()
        build(document, manifest, base)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)
        print(f"DOCX={output_path}")
    except (OSError, ValueError, json.JSONDecodeError, FileNotFoundError) as error:
        print(f"build_docx.py: {error}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
